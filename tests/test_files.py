"""
测试解析器 - 创建测试文件并测试解析
"""
import sys
from pathlib import Path

# 确保项目根目录在 Python 路径中
_project_root = Path(__file__).resolve().parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from parser.auto_detect import AutoDetectParser


def create_test_docx():
    """创建测试 DOCX 文件"""
    from docx import Document
    from docx.shared import Pt, Inches
    
    doc = Document()
    
    # 标题
    doc.add_heading('张三的简历', 0)
    
    # 个人信息
    doc.add_heading('个人信息', level=1)
    doc.add_paragraph('姓名：张三')
    doc.add_paragraph('电话：13800138000')
    doc.add_paragraph('邮箱：zhangsan@example.com')
    
    # 教育背景
    doc.add_heading('教育背景', level=1)
    doc.add_paragraph('北京大学 - 计算机科学 - 2018-2022')
    
    # 工作经历
    doc.add_heading('工作经历', level=1)
    doc.add_paragraph('百度 - 软件工程师 - 2022-至今')
    doc.add_paragraph('负责后端开发工作')
    
    # 保存
    test_dir = _project_root / "tests" / "testdata"
    test_dir.mkdir(exist_ok=True)
    docx_path = test_dir / "test_resume.docx"
    doc.save(str(docx_path))
    print(f"创建测试文件: {docx_path}")
    return docx_path


def test_docx_parser():
    """测试 DOCX 解析器"""
    print("\n" + "=" * 60)
    print("测试 DOCX 解析器")
    print("=" * 60)
    
    docx_path = create_test_docx()
    
    try:
        parser = AutoDetectParser()
        resume = parser.parse(docx_path)
        
        print(f"\n解析成功! 文件: {resume.file_name}")
        print(f"类型: {resume.file_type}")
        print(f"章节数: {len(resume.sections)}")
        print(f"\n章节内容:")
        for section in resume.sections:
            print(f"\n【{section.title}】")
            print(f"   {section.content[:100]}..." if len(section.content) > 100 else f"   {section.content}")
        
        return True
    except Exception as e:
        print(f"DOCX 解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_image_parser():
    """测试图片解析器"""
    print("\n" + "=" * 60)
    print("测试图片解析器")
    print("=" * 60)
    
    # 创建一个简单的测试图片
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        test_dir = _project_root / "tests" / "testdata"
        test_dir.mkdir(exist_ok=True)
        
        img_path = test_dir / "test_image.png"
        
        # 创建图片
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        # 添加文本
        text_lines = [
            "个人信息",
            "姓名：张三",
            "电话：13800138000",
            "",
            "教育背景",
            "北京大学",
            "",
            "工作经历",
            "百度公司 - 软件工程师",
        ]
        
        y = 50
        for line in text_lines:
            draw.text((50, y), line, fill='black')
            y += 40
        
        img.save(str(img_path))
        print(f"创建测试图片: {img_path}")
        
        # 解析
        parser = AutoDetectParser(default_lang="ch")
        resume = parser.parse(img_path)
        
        print(f"\n解析成功! 文件: {resume.file_name}")
        print(f"类型: {resume.file_type}")
        print(f"章节数: {len(resume.sections)}")
        print(f"\n章节内容:")
        for section in resume.sections:
            print(f"\n【{section.title}】")
            print(f"   {section.content[:100]}..." if len(section.content) > 100 else f"   {section.content}")
        
        return True
    except Exception as e:
        print(f"图片解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_pdf_parser():
    """测试 PDF 解析器"""
    print("\n" + "=" * 60)
    print("测试 PDF 解析器")
    print("=" * 60)
    
    # 使用已有的 a.pdf 文件
    pdf_path = Path("C:/Users/Just_me_cdd/Desktop/a.pdf")
    
    if not pdf_path.exists():
        print(f"PDF 文件不存在: {pdf_path}")
        return False
    
    try:
        parser = AutoDetectParser(default_mode="text")
        resume = parser.parse(pdf_path)
        
        print(f"\n解析成功! 文件: {resume.file_name}")
        print(f"类型: {resume.file_type}")
        print(f"章节数: {len(resume.sections)}")
        print(f"\n章节内容:")
        for section in resume.sections:
            print(f"\n【{section.title}】")
            print(f"   {section.content[:100]}..." if len(section.content) > 100 else f"   {section.content}")
        
        return True
    except Exception as e:
        print(f"PDF 解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("开始测试解析器...")
    
    results = {
        "pdf": test_pdf_parser(),
        "docx": test_docx_parser(),
        "image": test_image_parser(),
    }
    
    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    for name, success in results.items():
        status = "✅ 成功" if success else "❌ 失败"
        print(f"  {name.upper()}: {status}")


if __name__ == "__main__":
    main()