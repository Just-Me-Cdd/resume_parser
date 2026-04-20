"""
Word 简历解析器
"""

from pathlib import Path
from typing import Union, Optional, List, Dict
from docx import Document

from parser.base import BaseParser
from extractor.section_matcher import SectionMatcher
from extractor.cleaner import TextCleaner
from models.resume import Resume
from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """Word (DOCX) 简历解析器"""
    SUPPORTED_EXTENSIONS = [".docx"]
    
    def __init__(self):
        self.matcher = SectionMatcher()
        self.cleaner = TextCleaner()
    
    def parse(self, file_path: Union[str, Path]) -> Resume:
        path = self.validate_file(file_path)
        logger.info(f"开始解析 Word 简历: {path.name}")
        
        doc = Document(str(path))
        elements = self._extract_elements(doc)
        ocr_results = self._convert_to_ocr_format(elements)
        sections = self.matcher.match(ocr_results)
        raw_text = "\n".join([e["text"] for e in elements if e.get("text")])
        
        resume = Resume(sections=sections, raw_text=raw_text, file_name=path.name, file_type="docx")
        logger.info(f"解析完成，共 {len(sections)} 个章节")
        return resume
    
    def _extract_elements(self, doc: Document) -> List[Dict]:
        elements = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                elements.append({"type": "paragraph", "text": text, "style": para.style.name if para.style else "Normal"})
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    elements.append({"type": "table_row", "text": " | ".join(row_texts), "table_index": table_idx, "row_index": row_idx})
        
        return elements
    
    def _convert_to_ocr_format(self, elements: List[Dict]) -> List[Dict]:
        ocr_results = []
        y_offset = 0
        
        for element in elements:
            text = element["text"]
            cleaned_text = self.cleaner.clean(text)
            if not cleaned_text:
                continue
            
            if element["type"] == "paragraph":
                line_count = cleaned_text.count("\n") + 1
                height = 20 * line_count
                ocr_results.append({"text": cleaned_text, "bbox": [100, y_offset, 500, y_offset + height], "confidence": 1.0})
                y_offset += height + 10
            elif element["type"] == "table_row":
                ocr_results.append({"text": cleaned_text, "bbox": [50, y_offset, 550, y_offset + 20], "confidence": 0.95})
                y_offset += 25
        
        return ocr_results
    
    def _extract_tables(self, doc: Document) -> List[List[List[str]]]:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    @classmethod
    def can_parse(cls, file_path: Union[str, Path]) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS


class DOCParser(DOCXParser):
    """旧版 Word (DOC) 简历解析器"""
    SUPPORTED_EXTENSIONS = [".doc"]
    
    def __init__(self):
        super().__init__()
        logger.warning(".doc 格式支持有限，建议先将文档转换为 .docx 格式。")
    
    def parse(self, file_path: Union[str, Path]) -> Resume:
        path = self.validate_file(file_path)
        logger.info(f"开始解析 DOC 简历: {path.name}")
        
        try:
            return super().parse(file_path)
        except Exception as e:
            logger.error(f"DOC 解析失败: {e}")
            return self._parse_as_text(path)
    
    def _parse_as_text(self, file_path: Path) -> Resume:
        try:
            import subprocess
            result = subprocess.run(["antiword", str(file_path)], capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                text = result.stdout
                sections = self._simple分段(text)
                return Resume(sections=sections, raw_text=text, file_name=file_path.name, file_type="doc")
            else:
                raise RuntimeError("antiword 不可用")
        except Exception as e:
            logger.error(f"DOC 文本提取失败: {e}")
            return Resume(sections=[], raw_text="", file_name=file_path.name, file_type="doc")
    
    def _simple分段(self, text: str) -> List:
        lines = text.split("\n")
        sections = []
        current_title = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            is_title = False
            for pattern, name in self.matcher.compiled_patterns:
                if pattern.search(line):
                    is_title = True
                    break
            if is_title and len(line) < 50:
                if current_content:
                    sections.append({"title": current_title or "未分类", "content": "\n".join(current_content)})
                current_title = line
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections.append({"title": current_title or "未分类", "content": "\n".join(current_content)})
        
        return sections
